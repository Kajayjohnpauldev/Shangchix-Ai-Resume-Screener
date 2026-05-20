"""Resume text extraction.

Supports two upload formats:
  * PDF  → pypdf
  * Word → python-docx (.docx)

`extract_resume_text(filename, data)` is the router the API uses: it
dispatches on the file extension. Both extractors return an empty
string for unreadable input rather than raising, so a single bad upload
never crashes the batch — callers treat empty text as a zero-content
resume.
"""
from __future__ import annotations

import io
import re
from pathlib import Path

from pypdf import PdfReader
from pypdf.errors import PdfReadError


_WHITESPACE_RE = re.compile(r"\s+")

# Extensions we accept. PDF stays on pypdf; Word goes through python-docx.
PDF_EXTS = (".pdf",)
DOCX_EXTS = (".docx",)
SUPPORTED_EXTS = PDF_EXTS + DOCX_EXTS


def _clean(text: str) -> str:
    return _WHITESPACE_RE.sub(" ", text).strip()


def extract_text(pdf: str | Path | bytes) -> str:
    """Extract plain text from a PDF (path or raw bytes).

    Returns an empty string for unreadable or empty PDFs rather than
    raising.
    """
    try:
        if isinstance(pdf, (bytes, bytearray)):
            reader = PdfReader(io.BytesIO(pdf))
        else:
            reader = PdfReader(str(pdf))
    except (PdfReadError, FileNotFoundError, OSError):
        return ""

    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue

    return _clean("\n".join(pages))


def extract_docx(data: bytes) -> str:
    """Extract plain text from a .docx file given its raw bytes.

    Reads paragraphs and table cells (resumes often keep contact info or
    skills in tables). Returns an empty string for unreadable input.
    """
    try:
        from docx import Document  # imported lazily so PDF-only runs stay light
    except ImportError:
        return ""
    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        return ""

    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return _clean("\n".join(parts))


def extract_resume_text(filename: str, data: bytes) -> str:
    """Route to the right extractor based on the upload's file extension.

    Unknown extensions are attempted as PDF (the most common upload), so
    a mislabeled file still has a chance of parsing.
    """
    ext = Path(filename or "").suffix.lower()
    if ext in DOCX_EXTS:
        return extract_docx(data)
    return extract_text(data)


def candidate_name_from_filename(filename: str) -> str:
    """Turn 'Test_Candidate_One.pdf' (or .docx) into 'Test Candidate One'."""
    # Don't use Path.stem — it treats ".pdf" as a dotfile and returns
    # ".pdf" as the stem, which would render as ".Pdf".
    name = Path(filename).name
    for ext in (".pdf", ".docx", ".doc"):
        if name.lower().endswith(ext):
            name = name[: -len(ext)]
            break
    cleaned = re.sub(r"[_\-]+", " ", name).strip(" .")
    if not cleaned:
        return "Unknown Candidate"
    return cleaned.title()
